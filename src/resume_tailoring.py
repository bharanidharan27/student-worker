"""Create conservative tailored resume copies for saved jobs."""

from __future__ import annotations

import json
import re
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document

from src.eligibility.assessor import TECH_TERMS, _contains_term
from src.eligibility.models import EligibilityAssessment, JobRequirement
from src.eligibility.profile import ApplicantProfile, load_applicant_profile
from src.storage.db import DEFAULT_DB_PATH, get_job_by_id, insert_generated_document
from src.storage.models import GeneratedDocumentRecord
from src.utils.text_cleaner import normalize_whitespace


DEFAULT_EXTRACTED_RESUME_DIR = Path("resumes/extracted")
LEGACY_EXTRACTED_RESUME_DIR = Path("resume/extracted")
DEFAULT_TAILORED_RESUME_DIR = Path("resumes/tailored")

EXTRA_RESUME_TERMS = [
    "Blackboard",
    "Canvas",
    "LearningStudio",
    "Moodle",
    "WordPress",
    "PowerPoint",
    "Outlook",
]

LATEX_SKILL_BUCKETS = {
    "Platforms": {"Blackboard", "Canvas", "LearningStudio", "Moodle", "Open edX", "WordPress", "Zoom"},
    "Tools": {
        "AWS",
        "Docker",
        "Excel",
        "Git",
        "Google Workspace",
        "Jenkins",
        "Jira",
        "Microsoft Excel",
        "Microsoft Office",
        "Outlook",
        "Postman",
        "PowerPoint",
    },
    "Web & Systems": {"API", "APIs", "CSS", "HTML", "JavaScript", "React", "REST APIs", "TypeScript"},
    "Databases": {"MongoDB", "MySQL", "PostgreSQL", "Redis", "SQL"},
    "Support/Reporting": {"Power BI", "Tableau"},
}

MOJIBAKE_REPLACEMENTS = {
    "–": "--",
    "—": "---",
    "‘": "'",
    "’": "'",
    "“": '"',
    "”": '"',
    "\u00e2\u20ac\u201c": "--",
    "\u00e2\u20ac\u201d": "---",
    "\u00e2\u20ac\u02dc": "'",
    "\u00e2\u20ac\u2122": "'",
    "â€“": "--",
    "â€”": "---",
    "â€˜": "'",
    "â€™": "'",
    "â€œ": '"',
    "â€": '"',
    "Â": "",
}

GENERIC_TERM_WORDS = {
    "ability",
    "activities",
    "and",
    "assistant",
    "current",
    "experience",
    "knowledge",
    "minimum",
    "office",
    "preferred",
    "proficiency",
    "required",
    "requirements",
    "skills",
    "technology",
    "visual",
    "with",
    "work",
    "working",
}


@dataclass(frozen=True)
class TailoredResumeResult:
    job_id: int
    job_title: str
    source_resume_path: str
    output_resume_path: str
    output_dir: str
    notes_path: str
    generated_document_id: int
    additions: list[str]
    skipped: list[str]


@dataclass(frozen=True)
class ResumeEvidence:
    text: str
    label: str


@dataclass(frozen=True)
class ResumeSource:
    kind: str
    path: Path
    root_dir: Path | None = None


def tailor_resume_for_job(
    job_id: int,
    *,
    db_path: Path = DEFAULT_DB_PATH,
    extracted_dir: Path | None = None,
    output_root: Path = DEFAULT_TAILORED_RESUME_DIR,
    profile: ApplicantProfile | None = None,
) -> TailoredResumeResult:
    """Copy the recommended extracted resume source and add supported missing items."""

    row = get_job_by_id(job_id, db_path=db_path)
    if row is None:
        raise ValueError(f"No job found with local id {job_id}.")

    resolved_extracted_dir = _resolve_extracted_dir(extracted_dir)
    source_resume = _find_extracted_resume(row, resolved_extracted_dir)
    source_text = _read_resume_text(source_resume)
    applicant = profile or load_applicant_profile()
    evidence_bank = _evidence_bank(resolved_extracted_dir, applicant)
    assessment = _eligibility_from_row(row)
    additions, skipped = _supported_missing_additions(assessment, source_text, evidence_bank, applicant)

    output_dir = output_root / f"{job_id}-{_slug(row['title'] or 'job')}"
    output_dir.mkdir(parents=True, exist_ok=True)
    notes_path = output_dir / "tailoring_notes.md"
    output_resume = _copy_resume_source(source_resume, output_dir, job_id)
    if source_resume.kind == "latex":
        _sanitize_tex_source(output_resume)
    if additions:
        if source_resume.kind == "latex":
            additions, overflow = _apply_tailored_tex_additions(output_resume, additions)
            skipped.extend(
                [
                    f"{item} - not added to the LaTeX resume because it could not fit into the existing one-page skills layout."
                    for item in overflow
                ]
            )
        else:
            document = Document(str(output_resume))
            _append_tailored_docx_section(document, additions)
            document.save(str(output_resume))

    _write_notes(
        notes_path,
        job_id=job_id,
        job_title=row["title"] or "",
        source_resume=source_resume.path,
        output_resume=output_resume,
        additions=additions,
        skipped=skipped,
    )
    document_id = insert_generated_document(
        GeneratedDocumentRecord(
            job_id=job_id,
            document_type="resume",
            file_path=str(output_resume),
        ),
        db_path=db_path,
    )

    return TailoredResumeResult(
        job_id=job_id,
        job_title=row["title"] or "",
        source_resume_path=str(source_resume.path),
        output_resume_path=str(output_resume),
        output_dir=str(output_dir),
        notes_path=str(notes_path),
        generated_document_id=document_id,
        additions=additions,
        skipped=skipped,
    )


def _eligibility_from_row(row) -> EligibilityAssessment:
    raw_json = row["eligibility_json"]
    if not raw_json:
        raise ValueError("This job has no eligibility review yet. Review eligibility before tailoring a resume.")
    try:
        return EligibilityAssessment.model_validate(json.loads(raw_json))
    except (json.JSONDecodeError, ValueError, TypeError) as exc:
        raise ValueError("This job has an unreadable eligibility review. Re-review eligibility first.") from exc


def _resolve_extracted_dir(extracted_dir: Path | None) -> Path:
    if extracted_dir is not None:
        return extracted_dir
    if DEFAULT_EXTRACTED_RESUME_DIR.exists():
        return DEFAULT_EXTRACTED_RESUME_DIR
    return LEGACY_EXTRACTED_RESUME_DIR


def _find_extracted_resume(row, extracted_dir: Path) -> ResumeSource:
    resume_name = row["recommended_resume_name"]
    if not resume_name and row["recommended_resume_path"]:
        resume_name = Path(row["recommended_resume_path"]).name
    if not resume_name:
        raise ValueError("This job has no recommended resume. Re-score the job before tailoring.")

    target_stem = Path(resume_name).stem
    direct_latex_dir = extracted_dir / target_stem
    direct_latex_main = direct_latex_dir / "main.tex"
    if direct_latex_main.exists():
        return ResumeSource(kind="latex", path=direct_latex_main, root_dir=direct_latex_dir)

    direct_docx = extracted_dir / f"{target_stem}.docx"
    if direct_docx.exists():
        return ResumeSource(kind="docx", path=direct_docx)

    normalized_target = _normalize_filename_stem(target_stem)
    for candidate_dir in sorted(path for path in extracted_dir.iterdir() if path.is_dir()):
        candidate_main = candidate_dir / "main.tex"
        if candidate_main.exists() and _normalize_filename_stem(candidate_dir.name) == normalized_target:
            return ResumeSource(kind="latex", path=candidate_main, root_dir=candidate_dir)
    for candidate in extracted_dir.glob("*.docx"):
        if _normalize_filename_stem(candidate.stem) == normalized_target:
            return ResumeSource(kind="docx", path=candidate)
    raise FileNotFoundError(f"Could not find extracted LaTeX or DOCX source for {resume_name} in {extracted_dir}.")


def _evidence_bank(extracted_dir: Path, profile: ApplicantProfile) -> list[ResumeEvidence]:
    profile_text = " ".join(
        profile.technologies
        + profile.experience_domains
        + profile.certifications
        + profile.portfolio_links
        + profile.resume_keywords
    )
    evidence = [ResumeEvidence(profile_text, "applicant profile")]
    for tex_path in sorted(extracted_dir.glob("*/main.tex")):
        try:
            evidence.append(ResumeEvidence(_read_tex_text(tex_path), tex_path.parent.name))
        except Exception:
            continue
    for docx_path in sorted(extracted_dir.glob("*.docx")):
        try:
            evidence.append(ResumeEvidence(_read_docx_text(docx_path), docx_path.name))
        except Exception:
            continue
    return evidence


def _read_resume_text(source_resume: ResumeSource) -> str:
    if source_resume.kind == "latex":
        return _read_tex_text(source_resume.path)
    return _read_docx_text(source_resume.path)


def _read_docx_text(path: Path) -> str:
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return normalize_whitespace(" ".join(parts))


def _read_tex_text(path: Path) -> str:
    text = path.read_text(encoding="utf-8", errors="ignore")
    text = re.sub(r"(?<!\\)%.*", " ", text)
    text = (
        text.replace(r"\&", " & ")
        .replace(r"\%", " % ")
        .replace(r"\_", "_")
        .replace(r"\$", " $ ")
        .replace("\\\\", " ")
    )
    text = re.sub(r"\\[A-Za-z@]+(?:\[[^\]]*\])?", " ", text)
    text = text.replace("{", " ").replace("}", " ")
    text = text.replace("$", " ")
    return normalize_whitespace(text)


def _supported_missing_additions(
    assessment: EligibilityAssessment,
    source_text: str,
    evidence_bank: list[ResumeEvidence],
    profile: ApplicantProfile,
) -> tuple[list[str], list[str]]:
    additions: list[str] = []
    skipped: list[str] = []
    seen: set[str] = set()
    for requirement in assessment.requirements:
        if requirement.match != "missing":
            continue
        candidate = _addition_for_requirement(requirement, source_text, evidence_bank, profile)
        if candidate is None:
            skipped.append(f"{requirement.text} - no supporting evidence found in profile or extracted resumes.")
            continue
        statement, source_label = candidate
        key = statement.lower()
        if key in seen:
            continue
        seen.add(key)
        additions.append(f"{statement} Evidence: {source_label}.")
    return additions, skipped


def _addition_for_requirement(
    requirement: JobRequirement,
    source_text: str,
    evidence_bank: list[ResumeEvidence],
    profile: ApplicantProfile,
) -> tuple[str, str] | None:
    if requirement.category not in {"technology", "experience", "portfolio", "availability"}:
        return None

    terms = _candidate_terms(requirement)
    if requirement.category == "availability":
        if profile.available_hours_per_week is None:
            return None
        statement = f"Available up to {profile.available_hours_per_week} hours/week"
        if _contains_term(source_text, statement) or f"{profile.available_hours_per_week} hours/week" in source_text:
            return None
        return statement, "applicant profile"

    if requirement.category == "portfolio":
        if profile.portfolio_links and not _contains_term(source_text, "portfolio"):
            return "Portfolio available upon request", "applicant profile"
        return None

    supported_terms: list[str] = []
    source_labels: list[str] = []
    for term in terms:
        if _contains_term(source_text, term):
            continue
        evidence_label = _first_evidence([term], evidence_bank)
        if evidence_label:
            supported_terms.append(term)
            source_labels.append(evidence_label)

    supported_terms = _unique_preserve_order(supported_terms)
    source_labels = _unique_preserve_order(source_labels)
    if not supported_terms:
        return None

    if requirement.category == "technology":
        return f"Experience with {', '.join(supported_terms)}", ", ".join(source_labels)
    return f"Experience with {', '.join(supported_terms)}", ", ".join(source_labels)


def _candidate_terms(requirement: JobRequirement) -> list[str]:
    text = f"{requirement.text} {requirement.source_quote}"
    terms: list[str] = []
    for known_term in TECH_TERMS + EXTRA_RESUME_TERMS:
        if _contains_term(text, known_term):
            terms.append(known_term)

    for parenthetical in re.findall(r"\(([^)]+)\)", text):
        for term in re.split(r",|;|/|\bor\b", parenthetical):
            terms.append(term)

    for phrase in re.findall(r"\b[A-Z][A-Za-z0-9+#.-]*(?:\s+[A-Z][A-Za-z0-9+#.-]*){0,3}\b", text):
        terms.append(phrase)

    cleaned = [_clean_term(term) for term in terms]
    return [
        term
        for term in _unique_preserve_order(cleaned)
        if term and term.lower() not in GENERIC_TERM_WORDS and len(term) > 1
    ]


def _first_evidence(terms: list[str], evidence_bank: list[ResumeEvidence]) -> str | None:
    for evidence in evidence_bank:
        if all(_contains_term(evidence.text, term) for term in terms):
            return evidence.label
    return None


def _copy_resume_source(source_resume: ResumeSource, output_dir: Path, job_id: int) -> Path:
    if source_resume.kind == "latex":
        if source_resume.root_dir is None:
            raise ValueError("Latex resume source is missing its source directory.")
        shutil.copytree(source_resume.root_dir, output_dir, dirs_exist_ok=True)
        return output_dir / source_resume.path.name

    output_docx = output_dir / f"{source_resume.path.stem}_tailored_job_{job_id}.docx"
    shutil.copy2(source_resume.path, output_docx)
    return output_docx


def _append_tailored_docx_section(document, additions: list[str]) -> None:
    document.add_paragraph()
    document.add_heading("Targeted Skills", level=2)
    for addition in additions:
        try:
            paragraph = document.add_paragraph(style="List Bullet")
        except KeyError:
            paragraph = document.add_paragraph()
        paragraph.add_run(_resume_addition_text(addition))


def _apply_tailored_tex_additions(path: Path, additions: list[str]) -> tuple[list[str], list[str]]:
    content = path.read_text(encoding="utf-8")
    applied: list[str] = []
    overflow: list[str] = []
    updated = content
    for addition in additions:
        next_content, changed = _inline_tex_addition(updated, addition)
        if changed:
            updated = next_content
            applied.append(addition)
        else:
            overflow.append(addition)
    if updated != content:
        path.write_text(updated, encoding="utf-8")
    return applied, overflow


def _sanitize_tex_source(path: Path) -> None:
    content = path.read_text(encoding="utf-8")
    updated = content
    for source, replacement in MOJIBAKE_REPLACEMENTS.items():
        updated = updated.replace(source, replacement)
    updated = _replace_align_blocks(updated)
    updated = _wrap_availability_items(updated)
    updated = _close_open_resume_lists(updated)
    if updated != content:
        path.write_text(updated, encoding="utf-8")


def _replace_align_blocks(content: str) -> str:
    pattern = re.compile(r"\\begin\{align\}(.*?)\\end\{align\}", flags=re.DOTALL)

    def repl(match: re.Match[str]) -> str:
        body = re.sub(r"\s+", " ", match.group(1)).strip()
        return f"\\noindent\\small{{{body}}}\n"

    return pattern.sub(repl, content)


def _wrap_availability_items(content: str) -> str:
    pattern = re.compile(r"(\\section\{Availability\}\s*)(\\resumeItem\{.*?\})", flags=re.DOTALL)

    def repl(match: re.Match[str]) -> str:
        section_header = match.group(1).rstrip()
        item = match.group(2).strip()
        return f"{section_header}\n  \\resumeItemListStart\n    {item}\n  \\resumeItemListEnd\n"

    return pattern.sub(repl, content, count=1)


def _close_open_resume_lists(content: str) -> str:
    lines = content.splitlines()
    result: list[str] = []
    open_subheading_lists = 0

    for line in lines:
        stripped = line.strip()
        active = bool(stripped) and not stripped.startswith("%")
        should_close = active and (stripped.startswith(r"\section{") or stripped == r"\end{document}")
        if should_close and open_subheading_lists > 0:
            while open_subheading_lists > 0:
                result.append(r"\resumeSubHeadingListEnd")
                open_subheading_lists -= 1

        result.append(line)

        if not active:
            continue
        open_subheading_lists += stripped.count(r"\resumeSubHeadingListStart")
        open_subheading_lists -= stripped.count(r"\resumeSubHeadingListEnd")
        open_subheading_lists = max(0, open_subheading_lists)

    return "\n".join(result) + ("\n" if content.endswith("\n") else "")


def _inline_tex_addition(content: str, addition: str) -> tuple[str, bool]:
    addition_text = _resume_addition_text(addition)
    if addition_text.startswith("Available up to "):
        return _inline_availability_addition(content, addition_text)

    if addition_text == "Portfolio available upon request":
        return _inline_summary_addition(content, addition_text)

    if not addition_text.startswith("Experience with "):
        return content, False

    terms = [term.strip() for term in addition_text.removeprefix("Experience with ").split(",") if term.strip()]
    if not terms:
        return content, False

    updated = content
    for term in terms:
        bucket = _latex_skill_bucket(term)
        if bucket is None:
            return content, False
        updated, inserted = _append_term_to_skill_bucket(updated, bucket, term)
        if not inserted:
            return content, False
    return updated, True


def _inline_availability_addition(content: str, addition_text: str) -> tuple[str, bool]:
    pattern = re.compile(r"(?P<prefix>\\section\{Availability\}.*?\\resumeItem\{)(?P<body>.*?)(?P<suffix>\})", re.DOTALL)
    match = pattern.search(content)
    if match is None:
        return content, False

    body = match.group("body")
    if addition_text in body:
        return content, True

    updated_body = f"{body.rstrip()} $|$ {addition_text}"
    updated = f"{content[:match.start()]}{match.group('prefix')}{updated_body}{match.group('suffix')}{content[match.end():]}"
    return updated, True


def _inline_summary_addition(content: str, addition_text: str) -> tuple[str, bool]:
    pattern = re.compile(r"(?P<prefix>\\section\{Summary\}.*?\\noindent\\small\{)(?P<body>.*?)(?P<suffix>\})", re.DOTALL)
    match = pattern.search(content)
    if match is None:
        return content, False

    body = match.group("body")
    if addition_text in body:
        return content, True

    separator = " " if body.rstrip().endswith(".") else ". "
    updated_body = f"{body.rstrip()}{separator}{addition_text}."
    updated = f"{content[:match.start()]}{match.group('prefix')}{updated_body}{match.group('suffix')}{content[match.end():]}"
    return updated, True


def _latex_skill_bucket(term: str) -> str | None:
    for bucket, terms in LATEX_SKILL_BUCKETS.items():
        if term in terms:
            return bucket
    return None


def _append_term_to_skill_bucket(content: str, bucket: str, term: str) -> tuple[str, bool]:
    pattern = re.compile(rf"(?P<prefix>\\textbf\{{{re.escape(bucket)}\}}\{{:\s*)(?P<body>.*?)(?P<suffix>\}}\s*(?:\\\\)?)")
    match = pattern.search(content)
    if match is None:
        return content, False

    body = match.group("body")
    if _contains_term(body, term):
        return content, True

    delimiter = " $|$ " if "$|$" in body else ", "
    updated_body = f"{body.rstrip()}{delimiter}{term}"
    updated = f"{content[:match.start()]}{match.group('prefix')}{updated_body}{match.group('suffix')}{content[match.end():]}"
    return updated, True


def _write_notes(
    path: Path,
    *,
    job_id: int,
    job_title: str,
    source_resume: Path,
    output_resume: Path,
    additions: list[str],
    skipped: list[str],
) -> None:
    lines = [
        f"# Tailored Resume Notes - Job {job_id}",
        "",
        f"- Job: {job_title}",
        f"- Source resume: {source_resume}",
        f"- Tailored resume: {output_resume}",
        "",
        "## Added",
        "",
    ]
    lines.extend([f"- {addition}" for addition in additions] or ["- No supported missing items were added."])
    lines.extend(["", "## Skipped", ""])
    lines.extend([f"- {item}" for item in skipped] or ["- Nothing skipped."])
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def _clean_term(value: str) -> str:
    cleaned = re.sub(r"\b(?:e\.g\.?|i\.e\.?)\b,?", "", value, flags=re.IGNORECASE)
    cleaned = normalize_whitespace(cleaned).strip(" .:-")
    cleaned = re.sub(r"^(?:e\.g\.|including|such as)\s+", "", cleaned, flags=re.IGNORECASE)
    return cleaned.strip(" .:-")


def _normalize_filename_stem(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", value.lower())


def _slug(value: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")
    return slug[:80] or "job"


def _unique_preserve_order(values: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for value in values:
        key = value.lower()
        if key in seen:
            continue
        seen.add(key)
        result.append(value)
    return result


def _resume_addition_text(value: str) -> str:
    return re.sub(r"\s+Evidence:\s+.+$", "", value).strip()


def _latex_escape(value: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(char, char) for char in value)
