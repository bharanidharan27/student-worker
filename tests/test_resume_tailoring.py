from pathlib import Path

from docx import Document

from src.eligibility.models import EligibilityAssessment, JobRequirement
from src.resume_tailoring import tailor_resume_for_job
from src.storage.db import upsert_job
from src.storage.models import JobRecord


def _write_docx(path: Path, lines: list[str]) -> None:
    document = Document()
    for line in lines:
        document.add_paragraph(line)
    document.save(str(path))


def _read_docx(path: Path) -> str:
    document = Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _write_tex(path: Path, body_lines: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        "\n".join(
            [
                r"\documentclass{article}",
                r"\newcommand{\resumeItem}[1]{\item #1}",
                r"\newcommand{\resumeItemListStart}{\begin{itemize}}",
                r"\newcommand{\resumeItemListEnd}{\end{itemize}}",
                r"\begin{document}",
                *body_lines,
                r"\end{document}",
                "",
            ]
        ),
        encoding="utf-8",
    )


def test_tailor_resume_prefers_latex_source_and_adds_supported_missing_items(tmp_path: Path) -> None:
    db_path = tmp_path / "jobs.sqlite"
    extracted_dir = tmp_path / "resumes" / "extracted"
    output_root = tmp_path / "resumes" / "tailored"
    extracted_dir.mkdir(parents=True)
    _write_tex(extracted_dir / "Base_Resume" / "main.tex", [r"Python support experience."])
    _write_tex(extracted_dir / "Zoom_Resume" / "main.tex", [r"Zoom support and Canvas course shell experience."])
    assessment = EligibilityAssessment(
        status="needs_review",
        summary="Review gaps.",
        requirements=[
            JobRequirement(
                text="Knowledge or experience with Zoom.",
                priority="must",
                category="technology",
                source_quote="Working knowledge of Zoom is required.",
                confidence=0.9,
                match="missing",
            ),
            JobRequirement(
                text="Knowledge or experience with UnsupportedTool.",
                priority="preferred",
                category="technology",
                source_quote="UnsupportedTool preferred.",
                confidence=0.7,
                match="missing",
            ),
        ],
    )
    job_id = upsert_job(
        JobRecord(
            workday_id="JR-tailor",
            title="Technology Consultant",
            raw_description="Support Zoom.",
            recommended_resume_name="Base_Resume.pdf",
            recommended_resume_path="resumes/master/Base_Resume.pdf",
            eligibility_status="needs_review",
            eligibility_json=assessment.model_dump_json(),
        ),
        db_path=db_path,
    )

    result = tailor_resume_for_job(
        job_id,
        db_path=db_path,
        extracted_dir=extracted_dir,
        output_root=output_root,
    )

    output_path = Path(result.output_resume_path)
    assert output_path.exists()
    output_text = output_path.read_text(encoding="utf-8")
    assert output_path.name == "main.tex"
    assert "Targeted Skills" in output_text
    assert "Experience with Zoom" in output_text
    assert "Evidence:" not in output_text
    assert "UnsupportedTool" not in output_text
    assert "UnsupportedTool" in Path(result.notes_path).read_text(encoding="utf-8")


def test_tailor_resume_requires_eligibility_review(tmp_path: Path) -> None:
    db_path = tmp_path / "jobs.sqlite"
    extracted_dir = tmp_path / "resumes" / "extracted"
    extracted_dir.mkdir(parents=True)
    _write_docx(extracted_dir / "Base_Resume.docx", ["Python support experience."])
    job_id = upsert_job(
        JobRecord(
            workday_id="JR-no-review",
            title="Technology Consultant",
            raw_description="Support Zoom.",
            recommended_resume_name="Base_Resume.pdf",
            recommended_resume_path="resumes/master/Base_Resume.pdf",
        ),
        db_path=db_path,
    )

    try:
        tailor_resume_for_job(job_id, db_path=db_path, extracted_dir=extracted_dir)
    except ValueError as exc:
        assert "eligibility review" in str(exc)
    else:
        raise AssertionError("tailor_resume_for_job should require an eligibility review")


def test_tailor_resume_falls_back_to_docx_when_latex_source_is_missing(tmp_path: Path) -> None:
    db_path = tmp_path / "jobs.sqlite"
    extracted_dir = tmp_path / "resumes" / "extracted"
    output_root = tmp_path / "resumes" / "tailored"
    extracted_dir.mkdir(parents=True)
    _write_docx(extracted_dir / "Base_Resume.docx", ["Python support experience."])
    _write_docx(extracted_dir / "Zoom_Resume.docx", ["Zoom support and Canvas course shell experience."])
    assessment = EligibilityAssessment(
        status="needs_review",
        summary="Review gaps.",
        requirements=[
            JobRequirement(
                text="Knowledge or experience with Zoom.",
                priority="must",
                category="technology",
                source_quote="Working knowledge of Zoom is required.",
                confidence=0.9,
                match="missing",
            ),
        ],
    )
    job_id = upsert_job(
        JobRecord(
            workday_id="JR-tailor-docx",
            title="Technology Consultant",
            raw_description="Support Zoom.",
            recommended_resume_name="Base_Resume.pdf",
            recommended_resume_path="resumes/master/Base_Resume.pdf",
            eligibility_status="needs_review",
            eligibility_json=assessment.model_dump_json(),
        ),
        db_path=db_path,
    )

    result = tailor_resume_for_job(
        job_id,
        db_path=db_path,
        extracted_dir=extracted_dir,
        output_root=output_root,
    )

    output_path = Path(result.output_resume_path)
    output_text = _read_docx(output_path)
    assert output_path.suffix == ".docx"
    assert "Targeted Skills" in output_text
    assert "Experience with Zoom" in output_text
    assert "Evidence:" not in output_text
